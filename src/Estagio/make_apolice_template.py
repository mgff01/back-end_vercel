"""
Gera um template .docx simples para a Apolice de Seguro de Estagio, com tags
Jinja ({{ }}) limpas (python-docx escreve cada run inteiro, sem fragmentar).
Os ids das tags casam com CAMPOS_APOLICE em populate_db.py.
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

DEST = Path("media/modelos/modelo_apolice.docx")

doc = Document()

titulo = doc.add_paragraph()
titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = titulo.add_run("APÓLICE DE SEGURO DE ESTÁGIO")
run.bold = True
run.font.size = Pt(16)

doc.add_paragraph()  # espaco

linhas = [
    ("Seguradora:", "{{ seguradora }}"),
    ("Número da Apólice:", "{{ numero_apolice }}"),
    ("Segurado (Estagiário):", "{{ nome_aluno }}"),
    ("Matrícula:", "{{ matricula_aluno }}"),
    ("Empresa Concedente:", "{{ nome_empresa }}"),
    ("Capital Segurado (R$):", "{{ valor_segurado }}"),
    ("Início da Vigência:", "{{ vigencia_inicio }}"),
    ("Fim da Vigência:", "{{ vigencia_fim }}"),
]

tabela = doc.add_table(rows=0, cols=2)
tabela.style = "Table Grid"
for rotulo, tag in linhas:
    cells = tabela.add_row().cells
    r = cells[0].paragraphs[0].add_run(rotulo)
    r.bold = True
    cells[1].paragraphs[0].add_run(tag)

doc.add_paragraph()
nota = doc.add_paragraph()
nota.add_run(
    "Documento emitido conforme o inciso IV do Art. 9 da Lei 11.788/2008, "
    "garantindo a cobertura do(a) estagiário(a) contra acidentes pessoais "
    "durante o período de vigência do estágio."
).italic = True

DEST.parent.mkdir(parents=True, exist_ok=True)
doc.save(DEST)
print(f"Template criado: {DEST} ({DEST.stat().st_size} bytes)")

# Verificacao: docxtpl reconhece as tags e renderiza?
from docxtpl import DocxTemplate
import io

t = DocxTemplate(str(DEST))
varnames = sorted(t.get_undeclared_template_variables())
print("Tags reconhecidas pelo docxtpl:", varnames)
t.render({v: "TESTE" for v in varnames})
buf = io.BytesIO()
t.save(buf)
print(f"render() OK -> {buf.tell()} bytes")
